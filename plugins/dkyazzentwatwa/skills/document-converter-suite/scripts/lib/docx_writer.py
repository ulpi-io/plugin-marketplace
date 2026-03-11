from __future__ import annotations

from pathlib import Path
from typing import Iterable, List, Optional

from docx import Document


def write_docx_from_text(
    output_path: Path,
    title: Optional[str],
    paragraphs: Iterable[str],
    tables: Optional[List[List[List[str]]]] = None,
) -> None:
    """Write a simple DOCX with optional tables."""
    doc = Document()
    if title:
        doc.add_heading(title, level=1)

    for p in paragraphs:
        if p is None:
            continue
        txt = str(p).strip()
        if not txt:
            continue
        doc.add_paragraph(txt)

    if tables:
        for idx, t in enumerate(tables, start=1):
            # Add a little separation between the narrative text and tables.
            doc.add_paragraph("")
            doc.add_heading(f"Table {idx}", level=2)
            if not t:
                continue
            n_rows = len(t)
            n_cols = max((len(r) for r in t), default=0)
            if n_cols == 0:
                continue
            table = doc.add_table(rows=n_rows, cols=n_cols)
            for r_i, row in enumerate(t):
                for c_i in range(n_cols):
                    val = row[c_i] if c_i < len(row) else ""
                    table.cell(r_i, c_i).text = str(val or "")

    doc.save(str(output_path))
