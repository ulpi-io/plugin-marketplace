from __future__ import annotations

import shutil
import sys
from pathlib import Path
from typing import List, Optional

from docx import Document as DocxDocument

from .csv_reader import read_csv_content
from .csv_writer import write_csv_from_rows
from .docx_reader import read_docx_content
from .docx_writer import write_docx_from_text
from .html_reader import read_html_content
from .html_writer import write_html_from_sections
from .md_reader import read_md_content
from .md_writer import write_md_from_sections
from .pdf_reader import read_pdf_text
from .pdf_writer import write_pdf_from_sections
from .pptx_reader import read_pptx_content
from .pptx_writer import write_pptx_from_slides
from .txt_reader import read_txt_content
from .txt_writer import write_txt_from_lines
from .types import SUPPORTED_INPUT_EXTS, SUPPORTED_OUTPUT_EXTS
from .utils import chunk_list, guess_bullets, safe_cell_value
from .xlsx_reader import read_xlsx_content
from .xlsx_writer import write_xlsx_from_sheets


def _print_warnings(warnings: List[str]) -> None:
    """Print warnings to stderr with visual indicator."""
    for warning in warnings:
        print(f"⚠️  {warning}", file=sys.stderr)


def convert_document(
    input_path: Path,
    output_path: Path,
    max_pages: int = 200,
    max_chars: int = 300000,
    max_rows: int = 200,
    max_cols: int = 50,
    verbose: bool = False,
) -> None:
    """Convert a single document between PDF, DOCX, PPTX, XLSX.

    This is best-effort conversion focused on text and basic tables.
    """
    in_ext = input_path.suffix.lower()
    out_ext = output_path.suffix.lower()

    if in_ext not in SUPPORTED_INPUT_EXTS:
        raise ValueError(f"Unsupported input extension: {in_ext}")
    if out_ext not in SUPPORTED_OUTPUT_EXTS:
        raise ValueError(f"Unsupported output extension: {out_ext}")

    # Shortcut: same extension -> copy
    if in_ext == out_ext:
        if verbose:
            print(f"Copying (same format): {input_path} -> {output_path}")
        shutil.copy2(str(input_path), str(output_path))
        return

    # === Convert to DOCX ===
    if out_ext == ".docx":
        if in_ext == ".pdf":
            _pdf_to_docx(input_path, output_path, max_pages=max_pages, max_chars=max_chars)
            return
        if in_ext == ".pptx":
            _pptx_to_docx(input_path, output_path, max_chars=max_chars)
            return
        if in_ext == ".xlsx":
            _xlsx_to_docx(input_path, output_path, max_rows=max_rows, max_cols=max_cols)
            return
        if in_ext == ".md":
            _md_to_docx(input_path, output_path, max_chars=max_chars)
            return
        if in_ext == ".txt":
            _txt_to_docx(input_path, output_path, max_chars=max_chars)
            return
        if in_ext == ".csv":
            _csv_to_docx(input_path, output_path, max_rows=max_rows, max_cols=max_cols)
            return
        if in_ext == ".html":
            _html_to_docx(input_path, output_path, max_chars=max_chars)
            return

    # === Convert to PPTX ===
    if out_ext == ".pptx":
        if in_ext == ".pdf":
            _pdf_to_pptx(input_path, output_path, max_pages=max_pages, max_chars=max_chars)
            return
        if in_ext == ".docx":
            _docx_to_pptx(input_path, output_path, max_chars=max_chars)
            return
        if in_ext == ".xlsx":
            _xlsx_to_pptx(input_path, output_path, max_rows=max_rows, max_cols=max_cols)
            return
        if in_ext == ".md":
            _md_to_pptx(input_path, output_path, max_chars=max_chars)
            return
        if in_ext == ".txt":
            _txt_to_pptx(input_path, output_path, max_chars=max_chars)
            return
        if in_ext == ".csv":
            _csv_to_pptx(input_path, output_path, max_rows=max_rows, max_cols=max_cols)
            return
        if in_ext == ".html":
            _html_to_pptx(input_path, output_path, max_chars=max_chars)
            return

    # === Convert to XLSX ===
    if out_ext == ".xlsx":
        if in_ext == ".pdf":
            _pdf_to_xlsx(input_path, output_path, max_pages=max_pages, max_chars=max_chars)
            return
        if in_ext == ".docx":
            _docx_to_xlsx(input_path, output_path, max_chars=max_chars)
            return
        if in_ext == ".pptx":
            _pptx_to_xlsx(input_path, output_path, max_chars=max_chars)
            return
        if in_ext == ".md":
            _md_to_xlsx(input_path, output_path, max_chars=max_chars)
            return
        if in_ext == ".txt":
            _txt_to_xlsx(input_path, output_path, max_chars=max_chars)
            return
        if in_ext == ".csv":
            _csv_to_xlsx(input_path, output_path, max_rows=max_rows, max_cols=max_cols)
            return
        if in_ext == ".html":
            _html_to_xlsx(input_path, output_path, max_chars=max_chars)
            return

    # === Convert to PDF ===
    if out_ext == ".pdf":
        if in_ext == ".docx":
            _docx_to_pdf(input_path, output_path, max_chars=max_chars)
            return
        if in_ext == ".pptx":
            _pptx_to_pdf(input_path, output_path, max_chars=max_chars)
            return
        if in_ext == ".xlsx":
            _xlsx_to_pdf(input_path, output_path, max_rows=max_rows, max_cols=max_cols)
            return
        if in_ext == ".md":
            _md_to_pdf(input_path, output_path, max_chars=max_chars)
            return
        if in_ext == ".txt":
            _txt_to_pdf(input_path, output_path, max_chars=max_chars)
            return
        if in_ext == ".html":
            _html_to_pdf(input_path, output_path, max_chars=max_chars)
            return

    # === Convert to TXT ===
    if out_ext == ".txt":
        if in_ext == ".pdf":
            _pdf_to_txt(input_path, output_path, max_pages=max_pages, max_chars=max_chars)
            return
        if in_ext == ".docx":
            _docx_to_txt(input_path, output_path, max_chars=max_chars)
            return
        if in_ext == ".pptx":
            _pptx_to_txt(input_path, output_path, max_chars=max_chars)
            return
        if in_ext == ".xlsx":
            _xlsx_to_txt(input_path, output_path, max_rows=max_rows, max_cols=max_cols)
            return
        if in_ext == ".md":
            _md_to_txt(input_path, output_path, max_chars=max_chars)
            return
        if in_ext == ".csv":
            _csv_to_txt(input_path, output_path, max_rows=max_rows, max_cols=max_cols)
            return
        if in_ext == ".html":
            _html_to_txt(input_path, output_path, max_chars=max_chars)
            return

    # === Convert to CSV ===
    if out_ext == ".csv":
        if in_ext == ".xlsx":
            _xlsx_to_csv(input_path, output_path, max_rows=max_rows, max_cols=max_cols)
            return
        if in_ext == ".docx":
            _docx_to_csv(input_path, output_path, max_chars=max_chars)
            return
        if in_ext == ".pptx":
            _pptx_to_csv(input_path, output_path, max_chars=max_chars)
            return
        if in_ext == ".txt":
            _txt_to_csv(input_path, output_path, max_chars=max_chars)
            return
        if in_ext == ".html":
            _html_to_csv(input_path, output_path, max_chars=max_chars)
            return

    # === Convert to MD ===
    if out_ext == ".md":
        if in_ext == ".docx":
            _docx_to_md(input_path, output_path, max_chars=max_chars)
            return
        if in_ext == ".txt":
            _txt_to_md(input_path, output_path, max_chars=max_chars)
            return
        if in_ext == ".pdf":
            _pdf_to_md(input_path, output_path, max_pages=max_pages, max_chars=max_chars)
            return
        if in_ext == ".pptx":
            _pptx_to_md(input_path, output_path, max_chars=max_chars)
            return
        if in_ext == ".xlsx":
            _xlsx_to_md(input_path, output_path, max_rows=max_rows, max_cols=max_cols)
            return
        if in_ext == ".html":
            _html_to_md(input_path, output_path, max_chars=max_chars)
            return

    # === Convert to HTML ===
    if out_ext == ".html":
        if in_ext == ".md":
            _md_to_html(input_path, output_path, max_chars=max_chars)
            return
        if in_ext == ".docx":
            _docx_to_html(input_path, output_path, max_chars=max_chars)
            return
        if in_ext == ".txt":
            _txt_to_html(input_path, output_path, max_chars=max_chars)
            return
        if in_ext == ".pdf":
            _pdf_to_html(input_path, output_path, max_pages=max_pages, max_chars=max_chars)
            return
        if in_ext == ".pptx":
            _pptx_to_html(input_path, output_path, max_chars=max_chars)
            return
        if in_ext == ".xlsx":
            _xlsx_to_html(input_path, output_path, max_rows=max_rows, max_cols=max_cols)
            return
        if in_ext == ".csv":
            _csv_to_html(input_path, output_path, max_rows=max_rows, max_cols=max_cols)
            return

    raise ValueError(f"No conversion path implemented: {in_ext} -> {out_ext}")


# --------------------------
# PDF -> *
# --------------------------

def _pdf_to_docx(input_path: Path, output_path: Path, max_pages: int, max_chars: int) -> None:
    pdf = read_pdf_text(input_path, max_pages=max_pages, max_chars=max_chars)
    doc = DocxDocument()
    doc.add_heading(input_path.stem, level=1)
    for i, page_text in enumerate(pdf.pages, start=1):
        doc.add_heading(f"Page {i}", level=2)
        for ln in page_text.splitlines():
            ln = (ln or "").strip()
            if not ln:
                continue
            doc.add_paragraph(ln)
        if i < len(pdf.pages):
            doc.add_page_break()
    doc.save(str(output_path))


def _pdf_to_pptx(input_path: Path, output_path: Path, max_pages: int, max_chars: int) -> None:
    pdf = read_pdf_text(input_path, max_pages=max_pages, max_chars=max_chars)
    slides: List[dict] = []
    for i, page_text in enumerate(pdf.pages, start=1):
        slides.append({"title": f"Page {i}", "bullets": guess_bullets(page_text)})
    write_pptx_from_slides(output_path, slides, title=input_path.stem)


def _pdf_to_xlsx(input_path: Path, output_path: Path, max_pages: int, max_chars: int) -> None:
    pdf = read_pdf_text(input_path, max_pages=max_pages, max_chars=max_chars)
    rows: List[List[str]] = []
    for i, page_text in enumerate(pdf.pages, start=1):
        rows.append([f"=== Page {i} ==="])
        for ln in page_text.splitlines():
            ln = (ln or "").strip()
            if ln:
                rows.append([ln])
        rows.append([""])
    write_xlsx_from_sheets(output_path, [{"name": "PDF Text", "rows": rows}])


# --------------------------
# DOCX -> *
# --------------------------

def _docx_to_pptx(input_path: Path, output_path: Path, max_chars: int) -> None:
    content = read_docx_content(input_path, max_chars=max_chars)

    # Heuristic: if there are headings, start slides at headings.
    slides: List[dict] = []
    current_title: Optional[str] = None
    current_bullets: List[str] = []

    def flush():
        nonlocal current_title, current_bullets
        if current_title or current_bullets:
            slides.append({"title": current_title or "", "bullets": current_bullets})
        current_title = None
        current_bullets = []

    for p in content.paragraphs:
        # Treat lines that match known headings as slide boundaries
        if p in set(content.headings):
            flush()
            current_title = p
            continue
        # Chunk big docs into manageable slides
        current_bullets.append(p)
        if len(current_bullets) >= 10:
            flush()

    flush()
    if not slides:
        slides = [{"title": input_path.stem, "bullets": content.paragraphs[:10]}]

    # Normalize bullets a bit
    for s in slides:
        s["bullets"] = [b for b in guess_bullets("\n".join(s.get("bullets") or []), max_lines=12)]

    write_pptx_from_slides(output_path, slides, title=input_path.stem)


def _docx_to_xlsx(input_path: Path, output_path: Path, max_chars: int) -> None:
    content = read_docx_content(input_path, max_chars=max_chars)

    sheets: List[dict] = []

    # Text sheet
    text_rows = [[p] for p in content.paragraphs]
    sheets.append({"name": "Text", "rows": text_rows})

    # Tables -> their own sheets
    for idx, t in enumerate(content.tables, start=1):
        sheets.append({"name": f"Table{idx}", "rows": t.rows})

    write_xlsx_from_sheets(output_path, sheets)


def _docx_to_pdf(input_path: Path, output_path: Path, max_chars: int) -> None:
    content = read_docx_content(input_path, max_chars=max_chars)
    tables = [t.rows for t in content.tables]
    sections = [{"heading": input_path.stem, "paragraphs": content.paragraphs, "tables": tables}]
    write_pdf_from_sections(output_path, title=input_path.stem, sections=sections)


# --------------------------
# PPTX -> *
# --------------------------

def _pptx_to_docx(input_path: Path, output_path: Path, max_chars: int) -> None:
    content = read_pptx_content(input_path, max_chars=max_chars)

    doc = DocxDocument()
    doc.add_heading(input_path.stem, level=1)

    for s in content.slides:
        heading = s.title or f"Slide {s.index}"
        doc.add_heading(heading, level=2)

        # Text blocks -> bullets-ish
        for block in s.texts:
            for bullet in guess_bullets(block, max_lines=20):
                doc.add_paragraph(bullet, style="List Bullet")

        # Tables
        for t_i, t in enumerate(s.tables, start=1):
            if not t:
                continue
            doc.add_paragraph("")
            doc.add_heading(f"Table {t_i}", level=3)
            n_rows = len(t)
            n_cols = max((len(r) for r in t), default=0)
            if n_cols <= 0:
                continue
            table = doc.add_table(rows=n_rows, cols=n_cols)
            for r_i, row in enumerate(t):
                for c_i in range(n_cols):
                    table.cell(r_i, c_i).text = safe_cell_value(row[c_i] if c_i < len(row) else "")

    doc.save(str(output_path))


def _pptx_to_xlsx(input_path: Path, output_path: Path, max_chars: int) -> None:
    content = read_pptx_content(input_path, max_chars=max_chars)

    # Slides sheet: one row per slide with combined text
    rows: List[List[str]] = [["Slide", "Title", "Text"]]
    for s in content.slides:
        combined = "\n\n".join(s.texts)
        rows.append([str(s.index), s.title, combined])

    sheets: List[dict] = [{"name": "Slides", "rows": rows}]

    # Any tables: each table becomes its own sheet (bounded name)
    t_counter = 1
    for s in content.slides:
        for t in s.tables:
            name = f"Table{t_counter}"[:31]
            sheets.append({"name": name, "rows": t})
            t_counter += 1

    write_xlsx_from_sheets(output_path, sheets)


def _pptx_to_pdf(input_path: Path, output_path: Path, max_chars: int) -> None:
    content = read_pptx_content(input_path, max_chars=max_chars)
    sections: List[dict] = []
    for s in content.slides:
        heading = s.title or f"Slide {s.index}"
        paras: List[str] = []
        for block in s.texts:
            paras.extend(guess_bullets(block, max_lines=25))
        tables = s.tables
        sections.append({"heading": heading, "paragraphs": paras, "tables": tables})
    write_pdf_from_sections(output_path, title=input_path.stem, sections=sections)


# --------------------------
# XLSX -> *
# --------------------------

def _xlsx_to_docx(input_path: Path, output_path: Path, max_rows: int, max_cols: int) -> None:
    content = read_xlsx_content(input_path, max_rows=max_rows, max_cols=max_cols)
    _print_warnings(content.warnings)

    doc = DocxDocument()
    doc.add_heading(input_path.stem, level=1)

    for sh in content.sheets:
        doc.add_heading(sh.name, level=2)
        if not sh.cells:
            doc.add_paragraph("(empty sheet)")
            continue
        rows = sh.cells
        n_rows = len(rows)
        n_cols = max((len(r) for r in rows), default=0)
        table = doc.add_table(rows=n_rows, cols=n_cols)
        for r_i, row in enumerate(rows):
            for c_i in range(n_cols):
                table.cell(r_i, c_i).text = safe_cell_value(row[c_i] if c_i < len(row) else "")
        doc.add_paragraph("")

    doc.save(str(output_path))


def _xlsx_to_pptx(input_path: Path, output_path: Path, max_rows: int, max_cols: int) -> None:
    content = read_xlsx_content(input_path, max_rows=max_rows, max_cols=max_cols)
    _print_warnings(content.warnings)

    slides: List[dict] = []
    for sh in content.sheets:
        rows = sh.cells
        # If small enough, push it as a table; otherwise summarize first column as bullets.
        if len(rows) <= 20 and (max((len(r) for r in rows), default=0) <= 10):
            slides.append({"title": sh.name, "bullets": [], "tables": [rows]})
        else:
            first_col = [r[0] for r in rows if r and r[0]]
            bullets = first_col[:12] if first_col else ["(sheet too large to render as table)"]
            slides.append({"title": sh.name, "bullets": bullets})

    write_pptx_from_slides(output_path, slides, title=input_path.stem)


def _xlsx_to_pdf(input_path: Path, output_path: Path, max_rows: int, max_cols: int) -> None:
    content = read_xlsx_content(input_path, max_rows=max_rows, max_cols=max_cols)
    _print_warnings(content.warnings)
    sections: List[dict] = []
    for sh in content.sheets:
        sections.append({"heading": sh.name, "paragraphs": [], "tables": [sh.cells]})
    write_pdf_from_sections(output_path, title=input_path.stem, sections=sections)


# --------------------------
# TXT -> *
# --------------------------

def _txt_to_docx(input_path: Path, output_path: Path, max_chars: int) -> None:
    """Convert plain text to DOCX: lines become paragraphs."""
    content = read_txt_content(input_path, max_chars=max_chars)
    doc = DocxDocument()
    doc.add_heading(input_path.stem, level=1)
    for line in content.lines:
        if line.strip():
            doc.add_paragraph(line)
    doc.save(str(output_path))


def _txt_to_pptx(input_path: Path, output_path: Path, max_chars: int) -> None:
    """Convert plain text to PPTX: chunk lines into slides."""
    content = read_txt_content(input_path, max_chars=max_chars)
    slides: List[dict] = []

    # Chunk lines into slides (10-12 lines per slide)
    lines = [l for l in content.lines if l.strip()]
    for chunk in chunk_list(lines, size=12):
        slides.append({"title": "", "bullets": chunk})

    if not slides:
        slides = [{"title": input_path.stem, "bullets": ["(empty file)"]}]

    write_pptx_from_slides(output_path, slides, title=input_path.stem)


def _txt_to_xlsx(input_path: Path, output_path: Path, max_chars: int) -> None:
    """Convert plain text to XLSX: one line per row."""
    content = read_txt_content(input_path, max_chars=max_chars)
    rows: List[List[str]] = [[line] for line in content.lines]
    write_xlsx_from_sheets(output_path, [{"name": "Text", "rows": rows}])


def _txt_to_pdf(input_path: Path, output_path: Path, max_chars: int) -> None:
    """Convert plain text to PDF: lines become paragraphs."""
    content = read_txt_content(input_path, max_chars=max_chars)
    paragraphs = [line for line in content.lines if line.strip()]
    sections: List[dict] = [{"heading": input_path.stem, "paragraphs": paragraphs, "tables": []}]
    write_pdf_from_sections(output_path, title=input_path.stem, sections=sections)


def _txt_to_csv(input_path: Path, output_path: Path, max_chars: int) -> None:
    """Convert plain text to CSV: one line per row, single column."""
    content = read_txt_content(input_path, max_chars=max_chars)
    rows: List[List[str]] = [[line] for line in content.lines]
    write_csv_from_rows(output_path, rows=rows)


# --------------------------
# * -> TXT
# --------------------------

def _pdf_to_txt(input_path: Path, output_path: Path, max_pages: int, max_chars: int) -> None:
    """Convert PDF to plain text with page separators."""
    pdf = read_pdf_text(input_path, max_pages=max_pages, max_chars=max_chars)
    lines: List[str] = []
    for i, page_text in enumerate(pdf.pages, start=1):
        lines.append(f"=== Page {i} ===")
        lines.extend(page_text.splitlines())
        lines.append("")
    write_txt_from_lines(output_path, lines)


def _docx_to_txt(input_path: Path, output_path: Path, max_chars: int) -> None:
    """Convert DOCX to plain text: flatten paragraphs with heading markers."""
    content = read_docx_content(input_path, max_chars=max_chars)
    lines: List[str] = []
    heading_set = set(content.headings)

    for para in content.paragraphs:
        if para in heading_set:
            lines.append("")
            lines.append(f"## {para}")
            lines.append("")
        else:
            lines.append(para)

    write_txt_from_lines(output_path, lines)


def _pptx_to_txt(input_path: Path, output_path: Path, max_chars: int) -> None:
    """Convert PPTX to plain text: slide title + bullets."""
    content = read_pptx_content(input_path, max_chars=max_chars)
    lines: List[str] = []

    for slide in content.slides:
        if slide.title:
            lines.append(f"## {slide.title}")
            lines.append("")
        for text in slide.texts:
            lines.append(f"  - {text}")
        lines.append("")

    write_txt_from_lines(output_path, lines)


def _xlsx_to_txt(input_path: Path, output_path: Path, max_rows: int, max_cols: int) -> None:
    """Convert XLSX to plain text: sheet headers + cell values."""
    content = read_xlsx_content(input_path, max_rows=max_rows, max_cols=max_cols)
    _print_warnings(content.warnings)
    lines: List[str] = []

    for sh in content.sheets:
        lines.append(f"### {sh.name}")
        lines.append("")
        for row in sh.cells:
            line = " | ".join(safe_cell_value(cell) for cell in row)
            lines.append(line)
        lines.append("")

    write_txt_from_lines(output_path, lines)


def _csv_to_txt(input_path: Path, output_path: Path, max_rows: int, max_cols: int) -> None:
    """Convert CSV to plain text: pipe-separated values."""
    content = read_csv_content(input_path, max_rows=max_rows, max_cols=max_cols)
    lines: List[str] = []

    for row in content.rows:
        lines.append(" | ".join(row))

    write_txt_from_lines(output_path, lines)


# --------------------------
# CSV -> *
# --------------------------

def _csv_to_docx(input_path: Path, output_path: Path, max_rows: int, max_cols: int) -> None:
    """Convert CSV to DOCX: first row as heading, rest as table."""
    content = read_csv_content(input_path, max_rows=max_rows, max_cols=max_cols)
    doc = DocxDocument()
    doc.add_heading(input_path.stem, level=1)

    if not content.rows:
        doc.add_paragraph("(empty file)")
    else:
        # Create table from CSV rows
        n_rows = len(content.rows)
        n_cols = max((len(r) for r in content.rows), default=0)
        table = doc.add_table(rows=n_rows, cols=n_cols)

        for r_i, row in enumerate(content.rows):
            for c_i in range(n_cols):
                cell_value = row[c_i] if c_i < len(row) else ""
                table.cell(r_i, c_i).text = cell_value

    doc.save(str(output_path))


def _csv_to_pptx(input_path: Path, output_path: Path, max_rows: int, max_cols: int) -> None:
    """Convert CSV to PPTX: headers + rows as table slide."""
    content = read_csv_content(input_path, max_rows=max_rows, max_cols=max_cols)

    if not content.rows:
        slides = [{"title": input_path.stem, "bullets": ["(empty file)"]}]
    else:
        slides = [{"title": input_path.stem, "bullets": [], "tables": [content.rows]}]

    write_pptx_from_slides(output_path, slides, title=input_path.stem)


def _csv_to_xlsx(input_path: Path, output_path: Path, max_rows: int, max_cols: int) -> None:
    """Convert CSV to XLSX: direct mapping."""
    content = read_csv_content(input_path, max_rows=max_rows, max_cols=max_cols)
    write_xlsx_from_sheets(output_path, [{"name": "Sheet1", "rows": content.rows}])


# --------------------------
# * -> CSV
# --------------------------

def _xlsx_to_csv(input_path: Path, output_path: Path, max_rows: int, max_cols: int) -> None:
    """Convert XLSX to CSV: export first sheet only, warn if multiple sheets."""
    content = read_xlsx_content(input_path, max_rows=max_rows, max_cols=max_cols)
    _print_warnings(content.warnings)

    if len(content.sheets) > 1:
        print(f"⚠️  Warning: XLSX has {len(content.sheets)} sheets, exporting only first sheet '{content.sheets[0].name}'", file=sys.stderr)

    if content.sheets:
        write_csv_from_rows(output_path, rows=content.sheets[0].cells)
    else:
        write_csv_from_rows(output_path, rows=[])


def _docx_to_csv(input_path: Path, output_path: Path, max_chars: int) -> None:
    """Convert DOCX to CSV: extract tables only (text content lost)."""
    import sys
    content = read_docx_content(input_path, max_chars=max_chars)

    if not content.tables:
        print(f"⚠️  Warning: No tables found in DOCX, CSV will be empty", file=sys.stderr)
        write_csv_from_rows(output_path, rows=[])
    else:
        # Export first table only
        if len(content.tables) > 1:
            print(f"⚠️  Warning: DOCX has {len(content.tables)} tables, exporting only first table", file=sys.stderr)
        write_csv_from_rows(output_path, rows=content.tables[0].rows)


def _pptx_to_csv(input_path: Path, output_path: Path, max_chars: int) -> None:
    """Convert PPTX to CSV: extract tables only (text content lost)."""
    import sys
    content = read_pptx_content(input_path, max_chars=max_chars)

    # Collect all tables from all slides
    all_tables: List[List[List[str]]] = []
    for slide in content.slides:
        all_tables.extend(slide.tables)

    if not all_tables:
        print(f"⚠️  Warning: No tables found in PPTX, CSV will be empty", file=sys.stderr)
        write_csv_from_rows(output_path, rows=[])
    else:
        # Export first table only
        if len(all_tables) > 1:
            print(f"⚠️  Warning: PPTX has {len(all_tables)} tables, exporting only first table", file=sys.stderr)
        write_csv_from_rows(output_path, rows=all_tables[0])


# --------------------------
# MD -> *
# --------------------------

def _md_to_docx(input_path: Path, output_path: Path, max_chars: int) -> None:
    """Convert Markdown to DOCX with semantic mapping."""
    content = read_md_content(input_path, max_chars=max_chars)
    doc = DocxDocument()

    # Add title if present
    if content.title:
        doc.add_heading(content.title, level=1)

    for section in content.sections:
        section_type = section.get('type')

        if section_type == 'heading':
            level = min(section.get('level', 1), 9)  # DOCX supports up to level 9
            doc.add_heading(section.get('text', ''), level=level)

        elif section_type == 'paragraph':
            doc.add_paragraph(section.get('text', ''))

        elif section_type == 'list':
            items = section.get('items', [])
            ordered = section.get('ordered', False)
            style = 'List Number' if ordered else 'List Bullet'
            for item in items:
                doc.add_paragraph(item, style=style)

        elif section_type == 'table':
            headers = section.get('headers', [])
            rows = section.get('rows', [])
            if headers or rows:
                all_rows = [headers] + rows if headers else rows
                if all_rows:
                    n_rows = len(all_rows)
                    n_cols = max((len(r) for r in all_rows), default=0)
                    table = doc.add_table(rows=n_rows, cols=n_cols)
                    for r_i, row in enumerate(all_rows):
                        for c_i in range(n_cols):
                            cell_value = row[c_i] if c_i < len(row) else ""
                            table.cell(r_i, c_i).text = str(cell_value)

        elif section_type == 'code':
            code = section.get('code', '')
            doc.add_paragraph(code, style='No Spacing')

    doc.save(str(output_path))


def _md_to_pptx(input_path: Path, output_path: Path, max_chars: int) -> None:
    """Convert Markdown to PPTX: each H2 starts new slide."""
    content = read_md_content(input_path, max_chars=max_chars)
    slides: List[dict] = []

    current_title = content.title or input_path.stem
    current_bullets: List[str] = []
    current_tables: List[List[List[str]]] = []

    def flush_slide():
        nonlocal current_title, current_bullets, current_tables
        if current_title or current_bullets or current_tables:
            slides.append({
                "title": current_title or "",
                "bullets": current_bullets[:],
                "tables": current_tables[:]
            })
        current_title = ""
        current_bullets = []
        current_tables = []

    for section in content.sections:
        section_type = section.get('type')

        if section_type == 'heading':
            level = section.get('level', 1)
            text = section.get('text', '')

            # H2 and above start new slides
            if level >= 2:
                flush_slide()
                current_title = text
            else:
                # H1 becomes a bullet
                current_bullets.append(f"# {text}")

        elif section_type == 'paragraph':
            text = section.get('text', '')
            if text:
                current_bullets.append(text)

        elif section_type == 'list':
            items = section.get('items', [])
            current_bullets.extend(items)

        elif section_type == 'table':
            headers = section.get('headers', [])
            rows = section.get('rows', [])
            table_rows = [headers] + rows if headers else rows
            if table_rows:
                current_tables.append(table_rows)

        elif section_type == 'code':
            code = section.get('code', '')
            current_bullets.append(f"Code: {code[:100]}...")

    flush_slide()

    if not slides:
        slides = [{"title": content.title or input_path.stem, "bullets": ["(empty document)"]}]

    write_pptx_from_slides(output_path, slides, title=content.title or input_path.stem)


def _md_to_xlsx(input_path: Path, output_path: Path, max_chars: int) -> None:
    """Convert Markdown to XLSX: sections as sheets."""
    content = read_md_content(input_path, max_chars=max_chars)
    sheets: List[dict] = []

    current_sheet_name = "Content"
    current_rows: List[List[str]] = []

    for section in content.sections:
        section_type = section.get('type')

        if section_type == 'heading':
            level = section.get('level', 1)
            text = section.get('text', '')

            # H1 and H2 become sheet names
            if level <= 2 and current_rows:
                sheets.append({"name": current_sheet_name[:31], "rows": current_rows[:]})
                current_sheet_name = text[:31]
                current_rows = []
            else:
                current_rows.append([f"{'#' * level} {text}"])

        elif section_type == 'paragraph':
            text = section.get('text', '')
            if text:
                current_rows.append([text])

        elif section_type == 'list':
            items = section.get('items', [])
            for item in items:
                current_rows.append([f"- {item}"])

        elif section_type == 'table':
            headers = section.get('headers', [])
            rows = section.get('rows', [])
            if headers:
                current_rows.append(headers)
            current_rows.extend(rows)

    if current_rows:
        sheets.append({"name": current_sheet_name[:31], "rows": current_rows})

    if not sheets:
        sheets = [{"name": "Sheet1", "rows": [["(empty document)"]]}]

    write_xlsx_from_sheets(output_path, sheets)


def _md_to_pdf(input_path: Path, output_path: Path, max_chars: int) -> None:
    """Convert Markdown to PDF via sections."""
    content = read_md_content(input_path, max_chars=max_chars)
    sections: List[dict] = []

    current_heading = ""
    current_paragraphs: List[str] = []
    current_tables: List[List[List[str]]] = []

    def flush_section():
        nonlocal current_heading, current_paragraphs, current_tables
        if current_heading or current_paragraphs or current_tables:
            sections.append({
                "heading": current_heading,
                "paragraphs": current_paragraphs[:],
                "tables": current_tables[:]
            })
        current_heading = ""
        current_paragraphs = []
        current_tables = []

    for section in content.sections:
        section_type = section.get('type')

        if section_type == 'heading':
            flush_section()
            current_heading = section.get('text', '')

        elif section_type == 'paragraph':
            text = section.get('text', '')
            if text:
                current_paragraphs.append(text)

        elif section_type == 'list':
            items = section.get('items', [])
            for item in items:
                current_paragraphs.append(f"• {item}")

        elif section_type == 'table':
            headers = section.get('headers', [])
            rows = section.get('rows', [])
            table_rows = [headers] + rows if headers else rows
            if table_rows:
                current_tables.append(table_rows)

        elif section_type == 'code':
            code = section.get('code', '')
            current_paragraphs.append(f"[Code Block]\n{code}")

    flush_section()

    if not sections:
        sections = [{"heading": content.title or input_path.stem, "paragraphs": ["(empty document)"], "tables": []}]

    write_pdf_from_sections(output_path, title=content.title or input_path.stem, sections=sections)


def _md_to_txt(input_path: Path, output_path: Path, max_chars: int) -> None:
    """Convert Markdown to plain text: preserve markdown syntax."""
    content = read_md_content(input_path, max_chars=max_chars)
    lines: List[str] = []

    if content.title:
        lines.append(f"# {content.title}")
        lines.append("")

    for section in content.sections:
        section_type = section.get('type')

        if section_type == 'heading':
            level = section.get('level', 1)
            text = section.get('text', '')
            lines.append(f"{'#' * level} {text}")
            lines.append("")

        elif section_type == 'paragraph':
            text = section.get('text', '')
            if text:
                lines.append(text)
                lines.append("")

        elif section_type == 'list':
            items = section.get('items', [])
            ordered = section.get('ordered', False)
            for idx, item in enumerate(items, start=1):
                prefix = f"{idx}." if ordered else "-"
                lines.append(f"{prefix} {item}")
            lines.append("")

        elif section_type == 'table':
            headers = section.get('headers', [])
            rows = section.get('rows', [])
            if headers:
                lines.append('| ' + ' | '.join(headers) + ' |')
                lines.append('| ' + ' | '.join(['---'] * len(headers)) + ' |')
            for row in rows:
                padded = row + [''] * (len(headers) - len(row)) if headers else row
                lines.append('| ' + ' | '.join(str(c) for c in padded) + ' |')
            lines.append("")

        elif section_type == 'code':
            lang = section.get('language', '')
            code = section.get('code', '')
            lines.append(f"```{lang}")
            lines.append(code)
            lines.append("```")
            lines.append("")

    write_txt_from_lines(output_path, lines)


# --------------------------
# * -> MD
# --------------------------

def _docx_to_md(input_path: Path, output_path: Path, max_chars: int) -> None:
    """Convert DOCX to Markdown with structure preservation."""
    content = read_docx_content(input_path, max_chars=max_chars)
    sections: List[dict] = []
    heading_set = set(content.headings)

    for para in content.paragraphs:
        if para in heading_set:
            sections.append({'type': 'heading', 'level': 2, 'text': para})
        else:
            sections.append({'type': 'paragraph', 'text': para})

    # Add tables
    for table in content.tables:
        if table.rows:
            sections.append({
                'type': 'table',
                'headers': table.rows[0] if table.rows else [],
                'rows': table.rows[1:] if len(table.rows) > 1 else []
            })

    write_md_from_sections(output_path, sections, title=input_path.stem)


def _txt_to_md(input_path: Path, output_path: Path, max_chars: int) -> None:
    """Convert plain text to Markdown: lines become paragraphs."""
    content = read_txt_content(input_path, max_chars=max_chars)
    sections: List[dict] = []

    for line in content.lines:
        if line.strip():
            sections.append({'type': 'paragraph', 'text': line})

    write_md_from_sections(output_path, sections, title=input_path.stem)


def _pdf_to_md(input_path: Path, output_path: Path, max_pages: int, max_chars: int) -> None:
    """Convert PDF to Markdown: pages as sections."""
    pdf = read_pdf_text(input_path, max_pages=max_pages, max_chars=max_chars)
    sections: List[dict] = []

    for i, page_text in enumerate(pdf.pages, start=1):
        sections.append({'type': 'heading', 'level': 2, 'text': f"Page {i}"})
        for line in page_text.splitlines():
            if line.strip():
                sections.append({'type': 'paragraph', 'text': line})

    write_md_from_sections(output_path, sections, title=input_path.stem)


def _pptx_to_md(input_path: Path, output_path: Path, max_chars: int) -> None:
    """Convert PPTX to Markdown: slides as sections."""
    content = read_pptx_content(input_path, max_chars=max_chars)
    sections: List[dict] = []

    for slide in content.slides:
        if slide.title:
            sections.append({'type': 'heading', 'level': 2, 'text': slide.title})

        # Add text as list items
        if slide.texts:
            sections.append({'type': 'list', 'ordered': False, 'items': slide.texts})

        # Add tables
        for table in slide.tables:
            if table:
                sections.append({
                    'type': 'table',
                    'headers': table[0] if table else [],
                    'rows': table[1:] if len(table) > 1 else []
                })

    write_md_from_sections(output_path, sections, title=input_path.stem)


def _xlsx_to_md(input_path: Path, output_path: Path, max_rows: int, max_cols: int) -> None:
    """Convert XLSX to Markdown: sheets as sections with tables."""
    content = read_xlsx_content(input_path, max_rows=max_rows, max_cols=max_cols)
    _print_warnings(content.warnings)
    sections: List[dict] = []

    for sheet in content.sheets:
        sections.append({'type': 'heading', 'level': 2, 'text': sheet.name})

        if sheet.cells:
            sections.append({
                'type': 'table',
                'headers': sheet.cells[0] if sheet.cells else [],
                'rows': sheet.cells[1:] if len(sheet.cells) > 1 else []
            })

    write_md_from_sections(output_path, sections, title=input_path.stem)


# --------------------------
# HTML -> *
# --------------------------

def _html_to_docx(input_path: Path, output_path: Path, max_chars: int) -> None:
    """Convert HTML to DOCX: semantic structure preserved."""
    content = read_html_content(input_path, max_chars=max_chars)
    doc = DocxDocument()

    for section in content.sections:
        section_type = section.get('type')

        if section_type == 'heading':
            level = min(section.get('level', 1), 6)
            text = section.get('text', '')
            para = doc.add_paragraph(text, style=f'Heading {level}')

        elif section_type == 'paragraph':
            text = section.get('text', '')
            if text:
                doc.add_paragraph(text)

        elif section_type == 'list':
            items = section.get('items', [])
            for item in items:
                doc.add_paragraph(item, style='List Bullet')

        elif section_type == 'table':
            headers = section.get('headers', [])
            rows = section.get('rows', [])

            if headers or rows:
                num_cols = len(headers) if headers else (len(rows[0]) if rows else 1)
                num_rows = (1 if headers else 0) + len(rows)

                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Light Grid Accent 1'

                row_idx = 0
                if headers:
                    for col_idx, header in enumerate(headers):
                        table.cell(row_idx, col_idx).text = str(header)
                    row_idx += 1

                for row_data in rows:
                    for col_idx, cell_value in enumerate(row_data):
                        if col_idx < num_cols:
                            table.cell(row_idx, col_idx).text = str(cell_value)
                    row_idx += 1

        elif section_type == 'code':
            code = section.get('code', '')
            if code:
                para = doc.add_paragraph(code)
                para.style = 'No Spacing'

    doc.save(str(output_path))


def _html_to_pptx(input_path: Path, output_path: Path, max_chars: int) -> None:
    """Convert HTML to PPTX: headings as slide titles."""
    content = read_html_content(input_path, max_chars=max_chars)

    slides = []
    current_slide = {'title': content.title or input_path.stem, 'bullets': [], 'tables': []}

    for section in content.sections:
        section_type = section.get('type')

        if section_type == 'heading':
            # Start new slide on H2 or higher
            if section.get('level', 1) <= 2:
                if current_slide['bullets'] or current_slide['tables']:
                    slides.append(current_slide)
                current_slide = {'title': section.get('text', ''), 'bullets': [], 'tables': []}

        elif section_type == 'paragraph':
            text = section.get('text', '')
            if text:
                current_slide['bullets'].append(text)

        elif section_type == 'list':
            items = section.get('items', [])
            current_slide['bullets'].extend(items)

        elif section_type == 'table':
            headers = section.get('headers', [])
            rows = section.get('rows', [])
            if headers or rows:
                table_data = [headers] + rows if headers else rows
                current_slide['tables'].append(table_data)

    if current_slide['bullets'] or current_slide['tables']:
        slides.append(current_slide)

    if not slides:
        slides = [{'title': content.title or input_path.stem, 'bullets': ['(No content)'], 'tables': []}]

    write_pptx_from_slides(output_path, slides)


def _html_to_xlsx(input_path: Path, output_path: Path, max_chars: int) -> None:
    """Convert HTML to XLSX: tables to sheets."""
    content = read_html_content(input_path, max_chars=max_chars)

    sheets = []
    table_count = 0

    for section in content.sections:
        if section.get('type') == 'table':
            table_count += 1
            headers = section.get('headers', [])
            rows = section.get('rows', [])

            if headers or rows:
                cells = [headers] + rows if headers else rows
                sheets.append({
                    'name': f'Table{table_count}',
                    'cells': cells
                })

    if not sheets:
        # Create a sheet with document title
        sheets = [{
            'name': 'Content',
            'cells': [[content.title or input_path.stem]]
        }]

    write_xlsx_from_sheets(output_path, sheets)


def _html_to_pdf(input_path: Path, output_path: Path, max_chars: int) -> None:
    """Convert HTML to PDF via sections."""
    content = read_html_content(input_path, max_chars=max_chars)
    write_pdf_from_sections(output_path, content.sections, title=content.title)


def _html_to_txt(input_path: Path, output_path: Path, max_chars: int) -> None:
    """Convert HTML to plain text."""
    content = read_html_content(input_path, max_chars=max_chars)
    lines = []

    if content.title:
        lines.append(content.title)
        lines.append('=' * len(content.title))
        lines.append('')

    for section in content.sections:
        section_type = section.get('type')

        if section_type == 'heading':
            level = section.get('level', 1)
            text = section.get('text', '')
            prefix = '#' * level
            lines.append(f"{prefix} {text}")
            lines.append('')

        elif section_type == 'paragraph':
            text = section.get('text', '')
            if text:
                lines.append(text)
                lines.append('')

        elif section_type == 'list':
            items = section.get('items', [])
            for item in items:
                lines.append(f"• {item}")
            lines.append('')

        elif section_type == 'table':
            headers = section.get('headers', [])
            rows = section.get('rows', [])

            if headers:
                lines.append(' | '.join(headers))
                lines.append('-' * 40)

            for row in rows:
                lines.append(' | '.join(str(cell) for cell in row))

            lines.append('')

        elif section_type == 'code':
            code = section.get('code', '')
            lines.append(code)
            lines.append('')

    write_txt_from_lines(output_path, lines)


def _html_to_csv(input_path: Path, output_path: Path, max_chars: int) -> None:
    """Convert HTML to CSV: first table only."""
    content = read_html_content(input_path, max_chars=max_chars)

    for section in content.sections:
        if section.get('type') == 'table':
            headers = section.get('headers', [])
            rows = section.get('rows', [])

            write_csv_from_rows(output_path, rows, headers=headers if headers else None)
            return

    # No tables found, create a simple CSV
    write_csv_from_rows(output_path, [[content.title or input_path.stem]])


def _html_to_md(input_path: Path, output_path: Path, max_chars: int) -> None:
    """Convert HTML to Markdown: direct section mapping."""
    content = read_html_content(input_path, max_chars=max_chars)
    write_md_from_sections(output_path, content.sections, title=content.title)


# --------------------------
# * -> HTML
# --------------------------

def _md_to_html(input_path: Path, output_path: Path, max_chars: int) -> None:
    """Convert Markdown to HTML: direct section mapping."""
    content = read_md_content(input_path, max_chars=max_chars)
    write_html_from_sections(output_path, content.sections, title=content.title)


def _docx_to_html(input_path: Path, output_path: Path, max_chars: int) -> None:
    """Convert DOCX to HTML: paragraphs and tables."""
    content = read_docx_content(input_path, max_chars=max_chars)
    sections: List[dict] = []

    # Add headings
    for heading in content.headings:
        sections.append({'type': 'heading', 'level': 2, 'text': heading})

    # Add paragraphs
    for para in content.paragraphs:
        if para.strip():
            sections.append({'type': 'paragraph', 'text': para})

    # Add tables
    for table in content.tables:
        if table.cells:
            sections.append({
                'type': 'table',
                'headers': table.cells[0] if table.cells else [],
                'rows': table.cells[1:] if len(table.cells) > 1 else []
            })

    write_html_from_sections(output_path, sections, title=input_path.stem)


def _txt_to_html(input_path: Path, output_path: Path, max_chars: int) -> None:
    """Convert plain text to HTML: lines as paragraphs."""
    content = read_txt_content(input_path, max_chars=max_chars)
    sections: List[dict] = []

    for line in content.lines:
        if line.strip():
            sections.append({'type': 'paragraph', 'text': line})

    write_html_from_sections(output_path, sections, title=input_path.stem)


def _pdf_to_html(input_path: Path, output_path: Path, max_pages: int, max_chars: int) -> None:
    """Convert PDF to HTML: pages as sections."""
    pdf = read_pdf_text(input_path, max_pages=max_pages, max_chars=max_chars)
    sections: List[dict] = []

    for i, page_text in enumerate(pdf.pages, start=1):
        sections.append({'type': 'heading', 'level': 2, 'text': f"Page {i}"})
        for line in page_text.splitlines():
            if line.strip():
                sections.append({'type': 'paragraph', 'text': line})

    write_html_from_sections(output_path, sections, title=input_path.stem)


def _pptx_to_html(input_path: Path, output_path: Path, max_chars: int) -> None:
    """Convert PPTX to HTML: slides as sections."""
    content = read_pptx_content(input_path, max_chars=max_chars)
    sections: List[dict] = []

    for slide in content.slides:
        if slide.title:
            sections.append({'type': 'heading', 'level': 2, 'text': slide.title})

        # Add text as list items
        if slide.texts:
            sections.append({'type': 'list', 'ordered': False, 'items': slide.texts})

        # Add tables
        for table in slide.tables:
            if table:
                sections.append({
                    'type': 'table',
                    'headers': table[0] if table else [],
                    'rows': table[1:] if len(table) > 1 else []
                })

    write_html_from_sections(output_path, sections, title=input_path.stem)


def _xlsx_to_html(input_path: Path, output_path: Path, max_rows: int, max_cols: int) -> None:
    """Convert XLSX to HTML: sheets as sections with tables."""
    content = read_xlsx_content(input_path, max_rows=max_rows, max_cols=max_cols)
    _print_warnings(content.warnings)
    sections: List[dict] = []

    for sheet in content.sheets:
        sections.append({'type': 'heading', 'level': 2, 'text': sheet.name})

        if sheet.cells:
            sections.append({
                'type': 'table',
                'headers': sheet.cells[0] if sheet.cells else [],
                'rows': sheet.cells[1:] if len(sheet.cells) > 1 else []
            })

    write_html_from_sections(output_path, sections, title=input_path.stem)


def _csv_to_html(input_path: Path, output_path: Path, max_rows: int, max_cols: int) -> None:
    """Convert CSV to HTML: data as table."""
    content = read_csv_content(input_path, max_rows=max_rows, max_cols=max_cols)
    sections: List[dict] = []

    headers = content.headers if content.has_header else []
    rows = content.rows

    # If has_header, the first row is already in headers, so use remaining rows
    if content.has_header and rows:
        rows = rows

    sections.append({
        'type': 'table',
        'headers': headers,
        'rows': rows
    })

    write_html_from_sections(output_path, sections, title=input_path.stem)

