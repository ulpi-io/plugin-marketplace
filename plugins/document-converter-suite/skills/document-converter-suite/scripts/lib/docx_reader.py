from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document


@dataclass
class DocxTable:
    rows: List[List[str]]


@dataclass
class DocxContent:
    paragraphs: List[str]
    headings: List[str]
    heading_levels: Dict[str, int] = field(default_factory=dict)
    tables: List[DocxTable] = field(default_factory=list)


def _is_heading(paragraph) -> Optional[int]:
    """
    Detect if paragraph is a heading and return level (1-6) or None.

    Uses multi-heuristic approach:
    1. Primary: Style name contains "Heading"
    2. Secondary: Font size >= 14pt + bold
    3. Tertiary: Short text (< 100 chars) + ALL CAPS + bold
    """
    txt = (paragraph.text or "").strip()
    if not txt:
        return None

    # Primary: Style name
    try:
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.lower().startswith("heading"):
            # Extract level from style name (e.g., "Heading 1" -> 1)
            try:
                parts = style_name.split()
                if len(parts) >= 2 and parts[1].isdigit():
                    level = int(parts[1])
                    return min(max(level, 1), 6)  # Clamp to 1-6
            except Exception:
                pass
            return 2  # Default heading level if can't extract
    except Exception:
        pass

    # Secondary: Font size + bold
    try:
        if paragraph.runs:
            # Check first run for formatting
            run = paragraph.runs[0]
            font = run.font

            if font.size and font.bold:
                size_pt = font.size.pt
                if size_pt >= 18:
                    return 1
                elif size_pt >= 16:
                    return 2
                elif size_pt >= 14:
                    return 3
    except Exception:
        pass

    # Tertiary: Short + ALL CAPS + bold
    try:
        if len(txt) < 100 and txt.isupper() and paragraph.runs:
            run = paragraph.runs[0]
            if run.font.bold:
                return 2
    except Exception:
        pass

    return None


def read_docx_content(path: Path, max_chars: int = 300000) -> DocxContent:
    """Read text and basic tables from a .docx.

    Heuristics:
      - Heading detection: multi-heuristic (style, font size+bold, ALL CAPS+bold)
      - Tables: raw cell text
    """
    doc = Document(str(path))
    paras: List[str] = []
    headings: List[str] = []
    heading_levels: Dict[str, int] = {}

    budget = max_chars
    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if not txt:
            continue

        # Check if it's a heading
        heading_level = _is_heading(p)
        if heading_level is not None:
            headings.append(txt)
            heading_levels[txt] = heading_level

        if budget <= 0:
            break
        if len(txt) > budget:
            txt = txt[:budget]
        budget -= len(txt)
        paras.append(txt)

    tables: List[DocxTable] = []
    for t in doc.tables:
        rows: List[List[str]] = []
        for row in t.rows:
            rows.append([(cell.text or "").strip() for cell in row.cells])
        if rows:
            tables.append(DocxTable(rows=rows))

    return DocxContent(
        paragraphs=paras,
        headings=headings,
        heading_levels=heading_levels,
        tables=tables
    )
