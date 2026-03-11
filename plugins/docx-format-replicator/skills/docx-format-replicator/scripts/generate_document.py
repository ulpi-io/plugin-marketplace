#!/usr/bin/env python3
"""
Generate a Word document using extracted format information and content data.
"""

import sys
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def apply_style_from_config(paragraph, style_config):
    """Apply style configuration to a paragraph"""
    if not style_config:
        return

    # Apply fonts
    fonts = style_config.get('fonts', {})
    if fonts:
        for run in paragraph.runs:
            if 'ascii' in fonts and fonts['ascii']:
                run.font.name = fonts['ascii']
            if 'size' in fonts and fonts['size']:
                # Size is in half-points, convert to points
                run.font.size = Pt(int(fonts['size']) / 2)

    # Apply paragraph alignment
    para_props = style_config.get('paragraph', {})
    if para_props:
        alignment = para_props.get('alignment')
        if alignment == 'center':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == 'right':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif alignment == 'left':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alignment == 'both':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_heading_with_number(doc, text, level, number_text=None):
    """Add a heading with optional numbering"""
    heading = doc.add_heading(text, level=level)
    if number_text:
        # Prepend number to heading text
        heading.text = f"{number_text} {text}"
    return heading

def add_table_from_structure(doc, table_data, table_config):
    """Add a table based on configuration"""
    rows = table_data.get('rows', 1)
    cols = len(table_data.get('columns', [1]))

    table = doc.add_table(rows=rows, cols=cols)

    # Apply table style if available
    table_props = table_config.get('properties', {})
    style = table_props.get('style')
    if style:
        # python-docx uses style names, not IDs
        # You may need to map style IDs to names
        pass

    # Fill table with data if provided
    if 'cells' in table_data:
        for row_idx, row_data in enumerate(table_data['cells']):
            if row_idx < len(table.rows):
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx < len(table.columns):
                        table.rows[row_idx].cells[col_idx].text = str(cell_text)

    return table

def generate_document(format_path, content_path, output_path):
    """Generate a document from format and content configurations"""
    format_path = Path(format_path)
    content_path = Path(content_path)
    output_path = Path(output_path)

    if not format_path.exists():
        print(f"Error: Format file not found: {format_path}")
        return 1

    if not content_path.exists():
        print(f"Error: Content file not found: {content_path}")
        return 1

    # Load configurations
    with open(format_path, 'r', encoding='utf-8') as f:
        format_config = json.load(f)

    with open(content_path, 'r', encoding='utf-8') as f:
        content_data = json.load(f)

    # Create document
    doc = Document()

    # Set default style if available
    styles = format_config.get('styles', {})
    default_style = styles.get('1')  # Style ID 1 is typically "Normal"

    # Process content sections
    sections = content_data.get('sections', [])

    for section in sections:
        section_type = section.get('type')
        section_content = section.get('content', '')
        level = section.get('level', 1)
        number = section.get('number')

        if section_type == 'heading':
            add_heading_with_number(doc, section_content, level, number)

        elif section_type == 'paragraph':
            para = doc.add_paragraph(section_content)
            style_id = section.get('style_id')
            if style_id and style_id in styles:
                apply_style_from_config(para, styles[style_id])

        elif section_type == 'table':
            table_idx = section.get('table_index', 0)
            table_configs = format_config.get('tables', [])
            table_config = table_configs[table_idx] if table_idx < len(table_configs) else {}
            add_table_from_structure(doc, section, table_config)

        elif section_type == 'page_break':
            doc.add_page_break()

    # Save document
    doc.save(str(output_path))

    print(f"✅ Document generated successfully!")
    print(f"   Format: {format_path.name}")
    print(f"   Content: {content_path.name}")
    print(f"   Output: {output_path}")
    print(f"\nGenerated:")
    print(f"   - {len(sections)} sections")

    return 0

def main():
    if len(sys.argv) < 3:
        print("Usage: generate_document.py <format.json> <content.json> <output.docx>")
        print("\nGenerate a Word document from format and content configurations.")
        print("\nArguments:")
        print("  format.json  - Path to extracted format configuration")
        print("  content.json - Path to content data file")
        print("  output.docx  - Path for output Word document")
        return 1

    format_path = sys.argv[1]
    content_path = sys.argv[2]
    output_path = sys.argv[3] if len(sys.argv) > 3 else "output.docx"

    return generate_document(format_path, content_path, output_path)

if __name__ == '__main__':
    sys.exit(main())
