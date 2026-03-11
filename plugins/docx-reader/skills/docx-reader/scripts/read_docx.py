#!/usr/bin/env python3
"""
DOCX Reader Script
Extracts text content from Microsoft Word (.docx) files.
"""

from docx import Document
import sys
import os


def read_docx(file_path):
    """
    Read .docx file and extract text content.

    Args:
        file_path (str): Path to the .docx file

    Returns:
        str: Extracted text content
    """
    try:
        doc = Document(file_path)

        # Extract all paragraphs
        text_content = []
        for para in doc.paragraphs:
            if para.text.strip():  # Skip empty paragraphs
                text_content.append(para.text)

        # Extract tables if any
        for table in doc.tables:
            table_content = []
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    table_content.append(row_text)
            if table_content:
                text_content.append('\n'.join(table_content))

        return '\n\n'.join(text_content)

    except FileNotFoundError:
        return f"Error: File not found: {file_path}"
    except Exception as e:
        return f"Error reading file: {str(e)}"


def main():
    """Main entry point for the script."""
    if len(sys.argv) < 2:
        print("Usage: python read_docx.py <file_path>")
        sys.exit(1)

    file_path = sys.argv[1]

    if not os.path.exists(file_path):
        print(f"Error: File not found: {file_path}")
        sys.exit(1)

    if not file_path.lower().endswith('.docx'):
        print("Warning: File does not have .docx extension")

    content = read_docx(file_path)
    print(content)


if __name__ == "__main__":
    main()
