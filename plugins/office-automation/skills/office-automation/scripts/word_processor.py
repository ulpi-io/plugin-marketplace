#!/usr/bin/env python3
"""
Word 文档处理脚本
支持读取、写入、模板填充、表格提取等功能
"""

import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("错误：请安装 python-docx 库")
    print("运行：pip install python-docx")
    sys.exit(1)


def read_document(file_path: str) -> dict:
    """读取 Word 文档内容"""
    doc = Document(file_path)
    
    content = {
        "paragraphs": [],
        "tables": [],
        "sections": len(doc.sections)
    }
    
    # 读取段落
    for i, para in enumerate(doc.paragraphs):
        content["paragraphs"].append({
            "index": i,
            "text": para.text,
            "style": para.style.name if para.style else None
        })
    
    # 读取表格
    for i, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        content["tables"].append({
            "index": i,
            "rows": len(table.rows),
            "cols": len(table.columns),
            "data": table_data
        })
    
    return content


def write_document(file_path: str, content: str, title: str = None):
    """创建新的 Word 文档"""
    doc = Document()
    
    if title:
        doc.add_heading(title, 0)
    
    # 添加段落
    for paragraph in content.split('\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph)
    
    doc.save(file_path)
    print(f"✓ 文档已保存：{file_path}")


def fill_template(template_path: str, output_path: str, data: dict):
    """填充 Word 模板（替换 {{变量}} 格式）"""
    doc = Document(template_path)
    
    # 替换段落中的变量
    for para in doc.paragraphs:
        for key, value in data.items():
            if f"{{{{{key}}}}}" in para.text:
                para.text = para.text.replace(f"{{{{{key}}}}}", str(value))
    
    # 替换表格中的变量
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    if f"{{{{{key}}}}}" in cell.text:
                        cell.text = cell.text.replace(f"{{{{{key}}}}}", str(value))
    
    doc.save(output_path)
    print(f"✓ 模板已填充并保存：{output_path}")


def extract_tables(file_path: str, table_index: int = None, output_dir: str = "."):
    """提取 Word 文档中的表格到 CSV"""
    doc = Document(file_path)
    base_name = Path(file_path).stem
    
    tables_to_extract = []
    if table_index is not None:
        if table_index < len(doc.tables):
            tables_to_extract = [(table_index, doc.tables[table_index])]
        else:
            print(f"错误：表格索引 {table_index} 超出范围 (共 {len(doc.tables)} 个表格)")
            return
    else:
        tables_to_extract = list(enumerate(doc.tables))
    
    for idx, table in tables_to_extract:
        csv_path = Path(output_dir) / f"{base_name}_table_{idx}.csv"
        with open(csv_path, 'w', encoding='utf-8') as f:
            for row in table.rows:
                line = ','.join(f'"{cell.text}"' for cell in row.cells)
                f.write(line + '\n')
        print(f"✓ 表格 {idx} 已导出：{csv_path}")


def main():
    parser = argparse.ArgumentParser(description="Word 文档处理工具")
    subparsers = parser.add_subparsers(dest="command", help="命令")
    
    # read 命令
    read_parser = subparsers.add_parser("read", help="读取文档内容")
    read_parser.add_argument("file", help="Word 文件路径")
    
    # write 命令
    write_parser = subparsers.add_parser("write", help="创建新文档")
    write_parser.add_argument("output", help="输出文件路径")
    write_parser.add_argument("--content", required=True, help="文档内容")
    write_parser.add_argument("--title", help="文档标题")
    
    # template 命令
    template_parser = subparsers.add_parser("template", help="填充模板")
    template_parser.add_argument("template", help="模板文件路径")
    template_parser.add_argument("--output", required=True, help="输出文件路径")
    template_parser.add_argument("--data", required=True, help="JSON 格式数据")
    
    # extract 命令
    extract_parser = subparsers.add_parser("extract", help="提取表格")
    extract_parser.add_argument("file", help="Word 文件路径")
    extract_parser.add_argument("--table", type=int, help="表格索引（不指定则提取全部）")
    extract_parser.add_argument("--output-dir", default=".", help="输出目录")
    
    args = parser.parse_args()
    
    if args.command == "read":
        result = read_document(args.file)
        print(json.dumps(result, ensure_ascii=False, indent=2))
    
    elif args.command == "write":
        write_document(args.output, args.content, args.title)
    
    elif args.command == "template":
        data = json.loads(args.data)
        fill_template(args.template, args.output, data)
    
    elif args.command == "extract":
        extract_tables(args.file, args.table, args.output_dir)
    
    else:
        parser.print_help()


if __name__ == "__main__":
    main()
