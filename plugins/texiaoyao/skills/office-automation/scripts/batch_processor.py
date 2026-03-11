#!/usr/bin/env python3
"""
批量处理脚本
支持批量处理 Word 和 Excel 文件
"""

import argparse
import json
import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    import openpyxl
    import pandas as pd
except ImportError:
    print("错误：请安装必要的库")
    print("运行：pip install python-docx openpyxl pandas")
    sys.exit(1)


def batch_fill_templates(folder_path: str, template_path: str, data_file: str, output_folder: str):
    """批量填充 Word 模板"""
    folder = Path(folder_path)
    output = Path(output_folder)
    output.mkdir(parents=True, exist_ok=True)
    
    # 读取数据
    with open(data_file, 'r', encoding='utf-8') as f:
        data_list = json.load(f)
    
    template = Document(template_path)
    
    for i, data in enumerate(data_list):
        # 复制模板
        doc = Document(template_path)
        
        # 替换段落
        for para in doc.paragraphs:
            for key, value in data.items():
                if f"{{{{{key}}}}}" in para.text:
                    para.text = para.text.replace(f"{{{{{key}}}}}", str(value))
        
        # 替换表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data.items():
                        if f"{{{{{key}}}}}" in cell.text:
                            cell.text = cell.text.replace(f"{{{{{key}}}}}", str(value))
        
        # 保存
        output_name = data.get('filename', f"document_{i+1}.docx")
        output_path = output / output_name
        doc.save(output_path)
        print(f"✓ 已生成：{output_path}")
    
    print(f"\n✓ 批量处理完成：共生成 {len(data_list)} 个文件")


def batch_convert_folder(folder_path: str, from_format: str, to_format: str, output_folder: str):
    """批量转换文件格式"""
    folder = Path(folder_path)
    output = Path(output_folder)
    output.mkdir(parents=True, exist_ok=True)
    
    if from_format == "xlsx" and to_format == "csv":
        files = list(folder.glob("*.xlsx"))
        for file in files:
            df = pd.read_excel(file)
            csv_path = output / f"{file.stem}.csv"
            df.to_csv(csv_path, index=False, encoding='utf-8-sig')
            print(f"✓ 已转换：{file.name} → {csv_path.name}")
    
    elif from_format == "csv" and to_format == "xlsx":
        files = list(folder.glob("*.csv"))
        for file in files:
            df = pd.read_csv(file, encoding='utf-8')
            xlsx_path = output / f"{file.stem}.xlsx"
            df.to_excel(xlsx_path, index=False)
            print(f"✓ 已转换：{file.name} → {xlsx_path.name}")
    
    else:
        print(f"错误：不支持的转换 {from_format} → {to_format}")
        return
    
    print(f"\n✓ 批量转换完成：共处理 {len(files)} 个文件")


def batch_extract_tables(folder_path: str, output_folder: str):
    """批量提取 Word 文档中的表格"""
    folder = Path(folder_path)
    output = Path(output_folder)
    output.mkdir(parents=True, exist_ok=True)
    
    docx_files = list(folder.glob("*.docx"))
    
    for file in docx_files:
        doc = Document(file)
        base_name = file.stem
        
        for i, table in enumerate(doc.tables):
            csv_path = output / f"{base_name}_table_{i}.csv"
            with open(csv_path, 'w', encoding='utf-8') as f:
                for row in table.rows:
                    line = ','.join(f'"{cell.text}"' for cell in row.cells)
                    f.write(line + '\n')
            print(f"✓ 已提取：{file.name} → {csv_path.name}")
    
    print(f"\n✓ 批量提取完成：共处理 {len(docx_files)} 个文档")


def main():
    parser = argparse.ArgumentParser(description="批量处理工具")
    subparsers = parser.add_subparsers(dest="command", help="命令")
    
    # templates 命令
    tpl_parser = subparsers.add_parser("templates", help="批量填充模板")
    tpl_parser.add_argument("--folder", required=True, help="输出文件夹")
    tpl_parser.add_argument("--template", required=True, help="模板文件路径")
    tpl_parser.add_argument("--data", required=True, help="JSON 数据文件路径")
    
    # convert 命令
    conv_parser = subparsers.add_parser("convert", help="批量转换格式")
    conv_parser.add_argument("--folder", required=True, help="输入文件夹")
    conv_parser.add_argument("--from", dest="from_fmt", required=True, choices=["xlsx", "csv"])
    conv_parser.add_argument("--to", dest="to_fmt", required=True, choices=["xlsx", "csv"])
    conv_parser.add_argument("--output", required=True, help="输出文件夹")
    
    # extract 命令
    ext_parser = subparsers.add_parser("extract", help="批量提取表格")
    ext_parser.add_argument("--folder", required=True, help="包含 Word 文档的文件夹")
    ext_parser.add_argument("--output", required=True, help="输出文件夹")
    
    args = parser.parse_args()
    
    if args.command == "templates":
        batch_fill_templates(args.folder, args.template, args.data, args.output)
    
    elif args.command == "convert":
        batch_convert_folder(args.folder, args.from_fmt, args.to_fmt, args.output)
    
    elif args.command == "extract":
        batch_extract_tables(args.folder, args.output)
    
    else:
        parser.print_help()


if __name__ == "__main__":
    main()
